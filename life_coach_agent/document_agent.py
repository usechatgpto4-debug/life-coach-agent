"""
Document Creator Agent — Sub-agent for creating DOCX, PDF, XLSX documents.
Separated from the monolithic agent.py to follow multi-agent architecture.
"""

import os
import re
import sys
import json
import uuid
import logging
import subprocess
import tempfile
import traceback

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML
from openpyxl import Workbook
from openpyxl.styles import Font as XlFont, Alignment, PatternFill, Border, Side
import markdown

from google.adk.agents import Agent

logger = logging.getLogger(__name__)

# Directory to store generated files
EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)

# Font paths for Thai language support
_FONTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "fonts")
_THAI_FONT_REGULAR = os.path.join(_FONTS_DIR, "Sarabun-Regular.ttf")
_THAI_FONT_BOLD = os.path.join(_FONTS_DIR, "Sarabun-Bold.ttf")


# --- Helper: parse markdown-ish content into lines ---
def _parse_content_lines(content: str):
    """Parse content string into structured line tuples: (type, text)."""
    lines = content.split('\n')
    parsed = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            parsed.append(("blank", ""))
        elif stripped.startswith('## '):
            parsed.append(("h2", stripped[3:].strip()))
        elif stripped.startswith('### '):
            parsed.append(("h3", stripped[4:].strip()))
        elif stripped.startswith('- ') or stripped.startswith('• '):
            parsed.append(("bullet", stripped[2:].strip()))
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s*', '', stripped)
            parsed.append(("numbered", text))
        else:
            parsed.append(("text", stripped))
    return parsed


# ─── Tool: create_docx_document ────────────────────────────────────────────

def create_docx_document(title: str, content: str) -> dict:
    """Create a .docx document file that the user can download.

    Use this tool whenever the user asks to create a Word document or .docx file.

    Args:
        title: The document title (e.g., "คู่มือการสร้างหนังสือ", "แผนพัฒนาตนเอง")
        content: The full document content. Use '\\n' for line breaks and '## ' prefix for section headings.

    Returns:
        dict: A JSON object with type "docx_download" containing the file_id and filename for download.
    """
    doc = Document()

    # --- Style the document ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(14)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacer

    # Parse content into sections
    for line_type, text in _parse_content_lines(content):
        if line_type == "blank":
            doc.add_paragraph("")
        elif line_type == "h2":
            doc.add_heading(text, level=2)
        elif line_type == "h3":
            doc.add_heading(text, level=3)
        elif line_type == "bullet":
            doc.add_paragraph(text, style='List Bullet')
        elif line_type == "numbered":
            doc.add_paragraph(text, style='List Number')
        else:
            doc.add_paragraph(text)

    # Save file
    file_id = str(uuid.uuid4())
    safe_title = re.sub(r'[\\/:*?"<>|]', '', title).strip()
    filename = f"{safe_title}_{file_id[:8]}.docx"
    filepath = os.path.join(EXPORT_DIR, filename)
    doc.save(filepath)

    result = {
        "type": "docx_download",
        "file_id": file_id,
        "filename": filename,
        "title": title,
        "message": f"สร้างเอกสาร Word '{title}' เรียบร้อยแล้วครับ! คลิกปุ่มด้านล่างเพื่อดาวน์โหลด 📄"
    }

    return {"status": "success", "result": json.dumps(result, ensure_ascii=False)}


# ─── Tool: create_pdf_document ──────────────────────────────────────────────

def create_pdf_document(title: str, content: str) -> dict:
    """Create a .pdf document file that the user can download.

    Use this tool whenever the user asks to create a PDF file.

    Args:
        title: The document title (e.g., "คู่มือการสร้างหนังสือ", "แผนพัฒนาตนเอง")
        content: The full document content in Markdown format. Use standard markdown for headings, bold, lists, etc.

    Returns:
        dict: A JSON object with type "pdf_download" containing the file_id and filename for download.
    """
    file_id = str(uuid.uuid4())
    safe_title = re.sub(r'[\\/:*?"<>|]', '', title).strip()
    filename = f"{safe_title}_{file_id[:8]}.pdf"
    filepath = os.path.join(EXPORT_DIR, filename)

    # Convert markdown to HTML
    html_content = markdown.markdown(content)
    
    # Wrap in HTML template with styling for Thai fonts
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            @font-face {{
                font-family: 'Sarabun';
                src: url('file://{_THAI_FONT_REGULAR}') format('truetype');
                font-weight: normal;
            }}
            @font-face {{
                font-family: 'Sarabun';
                src: url('file://{_THAI_FONT_BOLD}') format('truetype');
                font-weight: bold;
            }}
            body {{
                font-family: 'Sarabun', sans-serif;
                font-size: 16pt;
                line-height: 1.5;
                margin: 2cm;
            }}
            h1, h2, h3, h4 {{
                font-weight: bold;
                color: #2c3e50;
            }}
            h1 {{ font-size: 24pt; text-align: center; margin-bottom: 30px; }}
            h2 {{ font-size: 20pt; margin-top: 20px; }}
            p {{ margin-bottom: 10px; word-break: break-word; }}
            ul, ol {{ margin-bottom: 10px; }}
            li {{ margin-bottom: 5px; }}
        </style>
    </head>
    <body>
        <h1>{title}</h1>
        {html_content}
    </body>
    </html>
    """

    # Generate PDF using WeasyPrint
    HTML(string=full_html).write_pdf(filepath)

    result = {
        "type": "pdf_download",
        "file_id": file_id,
        "filename": filename,
        "title": title,
        "message": f"สร้างเอกสาร PDF '{title}' เรียบร้อยแล้วครับ! คลิกปุ่มด้านล่างเพื่อดาวน์โหลด 📄"
    }

    return {"status": "success", "result": json.dumps(result, ensure_ascii=False)}


# ─── Tool: create_xlsx_document ─────────────────────────────────────────────

def create_xlsx_document(title: str, headers_json: str, rows_json: str) -> dict:
    """Create an .xlsx Excel spreadsheet file that the user can download.

    Use this tool whenever the user asks to create an Excel file, spreadsheet,
    table, checklist, or structured data export.

    Args:
        title: The spreadsheet title (e.g., "แผนการอ่าน 30 วัน", "Checklist พัฒนาตนเอง")
        headers_json: JSON array string of column headers, e.g. '["หัวข้อ", "รายละเอียด", "สถานะ"]'
        rows_json: JSON array of arrays string for row data, e.g. '[["อ่านหนังสือ", "30 นาที/วัน", "กำลังทำ"]]'

    Returns:
        dict: A JSON object with type "xlsx_download" containing the file_id and filename for download.
    """
    # Parse JSON string inputs
    try:
        headers = json.loads(headers_json) if isinstance(headers_json, str) else headers_json
    except (json.JSONDecodeError, TypeError):
        headers = ["Column 1"]

    try:
        rows = json.loads(rows_json) if isinstance(rows_json, str) else rows_json
    except (json.JSONDecodeError, TypeError):
        rows = []

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]  # Excel sheet name max 31 chars

    # --- Header styling ---
    header_font = XlFont(name='Arial', size=12, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='6C3CE0', end_color='6C3CE0', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin', color='D4D4D8'),
        right=Side(style='thin', color='D4D4D8'),
        top=Side(style='thin', color='D4D4D8'),
        bottom=Side(style='thin', color='D4D4D8'),
    )

    # --- Write title row ---
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(headers), 1))
    title_cell = ws.cell(row=1, column=1, value=title)
    title_cell.font = XlFont(name='Arial', size=16, bold=True, color='1A1A2E')
    title_cell.alignment = Alignment(horizontal='center', vertical='center')

    # --- Write headers (row 3) ---
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_idx, value=str(header))
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    # --- Write data rows ---
    data_font = XlFont(name='Arial', size=11)
    data_align = Alignment(vertical='center', wrap_text=True)
    alt_fill = PatternFill(start_color='F5F3FF', end_color='F5F3FF', fill_type='solid')

    for row_idx, row_data in enumerate(rows, 4):
        if not isinstance(row_data, (list, tuple)):
            row_data = [row_data]
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=str(value))
            cell.font = data_font
            cell.alignment = data_align
            cell.border = thin_border
            if (row_idx - 4) % 2 == 1:
                cell.fill = alt_fill

    # --- Auto-fit column widths (approximate) ---
    for col_idx in range(1, len(headers) + 1):
        max_len = len(str(headers[col_idx - 1])) if col_idx <= len(headers) else 10
        for row_data in rows:
            if isinstance(row_data, (list, tuple)) and col_idx - 1 < len(row_data):
                max_len = max(max_len, len(str(row_data[col_idx - 1])))
        col_letter = chr(64 + col_idx) if col_idx <= 26 else 'A'
        ws.column_dimensions[col_letter].width = min(max_len + 4, 50)

    # --- Save ---
    file_id = str(uuid.uuid4())
    safe_title = re.sub(r'[\\/:*?"<>|]', '', title).strip()
    filename = f"{safe_title}_{file_id[:8]}.xlsx"
    filepath = os.path.join(EXPORT_DIR, filename)
    wb.save(filepath)

    result = {
        "type": "xlsx_download",
        "file_id": file_id,
        "filename": filename,
        "title": title,
        "row_count": len(rows),
        "message": f"สร้างไฟล์ Excel '{title}' ({len(rows)} แถว) เรียบร้อยแล้วครับ! คลิกปุ่มด้านล่างเพื่อดาวน์โหลด 📊"
    }

    return {"status": "success", "result": json.dumps(result, ensure_ascii=False)}


# ─── Tool: execute_python_code ──────────────────────────────────────────────

def execute_python_code(code: str) -> dict:
    """Execute Python code and return the output.
    
    This is useful for generating custom files like PDF using fpdf, data analysis, 
    or any task that requires executing dynamic Python scripts.
    
    Args:
        code: The Python code to execute.
        
    Returns:
        dict: A JSON object containing the execution status and output or error.
    """
    try:
        # Create a temporary file to hold the code
        with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False) as temp_file:
            temp_file.write(code)
            temp_file_path = temp_file.name
            
        # Execute the code
        result = subprocess.run(
            [sys.executable, temp_file_path],
            capture_output=True,
            text=True,
            check=False,
            timeout=30,
            cwd=os.path.dirname(os.path.dirname(__file__))
        )
        
        # Clean up
        try:
            os.remove(temp_file_path)
        except Exception:
            pass
            
        if result.returncode == 0:
            return {"status": "success", "output": result.stdout}
        else:
            return {"status": "error", "error": result.stderr, "output": result.stdout}
            
    except subprocess.TimeoutExpired:
        return {"status": "error", "error": "Execution timed out after 30 seconds."}
    except Exception as e:
        logger.error("execute_python_code failed: %s", e, exc_info=True)
        return {"status": "error", "error": str(e), "traceback": traceback.format_exc()}


# ─── Document Agent Definition ──────────────────────────────────────────────

DOCUMENT_INSTRUCTION = """คุณคือ Document Creator — ผู้เชี่ยวชาญด้านการสร้างเอกสาร

คุณมีเครื่องมือ 4 ตัว:
1. create_docx_document — สร้างเอกสาร Word (.docx)
2. create_pdf_document — สร้างเอกสาร PDF
3. create_xlsx_document — สร้างไฟล์ Excel (.xlsx)
4. execute_python_code — รันโค้ด Python สำหรับงานพิเศษ

[กฎการทำงาน]
- สร้างเอกสารที่มีโครงสร้างชัดเจน มีหัวข้อ รายการ และเนื้อหาครบถ้วน
- ใช้ภาษาที่สุภาพ เป็นมืออาชีพ
- เมื่อสร้างเอกสารเสร็จ ให้ส่ง JSON result กลับ (ห้ามพิมพ์ข้อความอื่นนอก JSON)
- สำหรับ PDF: ใส่ Markdown ใน content ได้เลย เครื่องมือจะจัดการฟอนต์ภาษาไทยให้
- สำหรับ Excel: ต้องส่ง headers_json และ rows_json เป็น JSON array string
"""

document_agent = Agent(
    name="DocumentCreator",
    model="gemini-3.1-pro-preview",
    description="สร้างเอกสาร DOCX, PDF, XLSX และรันโค้ด Python — ใช้สำหรับงานที่ต้องการสร้างไฟล์เอกสาร ตาราง รายงาน หรือรันสคริปต์",
    instruction=DOCUMENT_INSTRUCTION,
    tools=[create_docx_document, create_pdf_document, create_xlsx_document, execute_python_code],
)
