"""
Bestseller Novel Architect — Automated Novel Generator Module
Adapted from the bestseller-novel-architect skill framework.
Generates full novels using a multi-step Gemini API pipeline.
"""

import os
import sys
import re
import json
import time
import uuid
import logging
import tempfile
from pydantic import BaseModel, Field
from typing import List, Optional, Callable

from google import genai
from google.genai import types
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

logger = logging.getLogger(__name__)

EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)

# --- Pydantic Data Models ---

class Character(BaseModel):
    """Three-Layer Character Model (Surface → Private → Blind Spot)."""
    name: str = Field(description="ชื่อตัวละคร")
    goal: str = Field(description="เป้าหมายภายนอกที่ชัดเจน")
    flaw_or_trauma: str = Field(description="ความขัดแย้งภายใน หรือปมบาดแผลในอดีต")
    arc_type: str = Field(description="พัฒนาการของตัวละคร (เชิงบวก, เชิงลบ, หรือซับซ้อน)")
    # Three-Layer Character Design
    surface_persona: str = Field(default="", description="สิ่งที่คนอื่นเห็น — บุคลิก ท่าทาง การแสดงออกภายนอก")
    private_self: str = Field(default="", description="สิ่งที่ซ่อนอยู่ — ความกลัว ความปรารถนาลึกๆ นิสัยส่วนตัว")
    blind_spot: str = Field(default="", description="สิ่งที่ตัวละครไม่รู้เกี่ยวกับตัวเอง — อคติ รูปแบบพฤติกรรมที่ซ้ำซาก")

class OutlineAct(BaseModel):
    act_number: int
    description: str
    key_plot_points: List[str]

class NovelOutline(BaseModel):
    title: str = Field(description="ชื่อเรื่องที่สั้น กระชับ 1-5 คำ สร้าง curiosity gap ไม่ใช้คำ AI ซ้ำซาก (เช่น shadows, echoes, whispers, fate, destiny) ต้องบ่งบอกแนวเรื่อง สร้างคำถามในใจผู้อ่าน")
    genre: str
    premise: str
    characters: List[Character]
    acts: List[OutlineAct]
    total_chapters: int

class SceneOutline(BaseModel):
    scene_number: int = Field(description="ลำดับของฉาก")
    setting: str = Field(description="สถานที่และเวลา")
    action_or_conflict: str = Field(description="เหตุการณ์หลักหรือความขัดแย้งในฉากนี้")

class ChapterOutline(BaseModel):
    chapter_number: int
    title: str = Field(description="ชื่อบทที่น่าดึงดูดใจ")
    scenes: List[SceneOutline] = Field(description="รายการฉากในบทนี้ (3-4 ฉาก)")


# --- Master System Prompt (9 Pillars + Anti-AI + Emotional Arc) ---

MASTER_PROMPT = """คุณคือนักเขียนนิยายระดับเบสต์เซลเลอร์และเจ้าของรางวัลโนเบล
หน้าที่ของคุณคือการเขียนนิยายโดยยึดหลัก 8 เสาหลักดังนี้:
1. โครงสร้าง 3 องก์ (Setup 25%, Confrontation 50%, Resolution 25%)
2. มุมมองการเล่าเรื่องแบบลึกซึ้ง (Deep POV)
3. ตัวละครต้องมีเป้าหมายภายนอกและบาดแผลทางจิตใจ (ทฤษฎีภูเขาน้ำแข็ง)
4. จังหวะการเล่าเรื่อง (Pacing) กระชับ มีเป้าหมายย่อยทุกบท
5. พัฒนาการของตัวละคร (Character Arc) ชัดเจน
6. โครงสร้างบท: เปิดฉากท่ามกลางเหตุการณ์ (In medias res) มีความขัดแย้งย่อย
7. ทุกบทต้องจบด้วย Hook (แอ็กชัน, การเปิดเผยความลับ, คำถาม, อารมณ์, Shift, Promise)
8. Show Don't Tell: เน้นประสาทสัมผัสทั้ง 5 ห้ามสรุปอารมณ์ตรงๆ

[THREE-LAYER CHARACTER — บังคับใช้ทุกฉาก]
- Surface Layer: แสดงบุคลิกภายนอกของตัวละครผ่านบทสนทนาและท่าทาง
- Private Layer: เผยความกลัว/ความปรารถนาผ่านความคิดภายในและนิสัยลับ
- Blind Spot Layer: ฝังสิ่งที่ตัวละครไม่รู้เกี่ยวกับตัวเอง ให้ผู้อ่านเห็นก่อนตัวละคร
- บทที่ 1-30%: เน้น Surface  |  30-60%: เปิด Private  |  60-100%: เผชิญ Blind Spot

[EMOTIONAL ARC — บังคับใช้ทุกบท]
1. วินิจฉัยอารมณ์ผู้อ่านขาเข้า (จากบทก่อนหน้า)
2. กำหนดอารมณ์ปลายทางของบทนี้
3. สร้างเส้นทางอารมณ์ที่ราบรื่น (ห้ามกระโดดข้ามอารมณ์)
4. วางจุดพีค (Peak Moment) ตรงกลาง-หลังของบท
5. ออกแบบจุดจบเพื่อส่งต่อไปยังบทถัดไป

[ANTI-AI PROSE — ข้อห้ามเด็ดขาด]
ห้ามใช้คำและรูปแบบเหล่านี้ในเนื้อหาโดยเด็ดขาด:
- คำ: leverage, utilize, robust, seamless, cutting-edge, innovative, facilitate, streamline, navigate, pivotal, embark, foster, harness, paradigm, holistic, delve, tapestry, symphony, testament
- รูปแบบ: "At its core", "In a world where", "It's important to note", "Let's explore", "What this means is"
- ประเภทประโยค: Hedging ("It's worth noting"), Hollow intensifiers ("Very", "Extremely"), Generic conclusions ("In conclusion")
- ห้ามใช้ em-dash (—) ซ้ำเกินกว่า 1 ครั้งต่อหน้า
- ห้ามสรุปซ้ำสิ่งที่เล่าไปแล้วในตอนจบบท
- ห้ามใช้ประโยคที่มีเพียงเพื่อบอกอารมณ์ผู้อ่าน (เช่น "This was truly heartbreaking")

[PROSE QUALITY — คุณภาพงานเขียน]
- ใช้คำกริยาที่ทรงพลังและเฉพาะเจาะจง แทนคำกริยาทั่วไป+คำวิเศษณ์
- สลับความยาวประโยค — ประโยคสั้นสร้าง impact ประโยคยาวสร้างจังหวะ
- เปิดย่อหน้าด้วยเนื้อหา ไม่ใช่ transition words (Moreover, Furthermore, Additionally)
- จบบทอย่างสะอาด ไม่ต้องสรุปซ้ำ thesis
- ทุกย่อหน้าต้องขับเคลื่อนความหมายไปข้างหน้า

[สำคัญ] ห้ามใส่ 'สารบัญ' หรือ 'Table of Contents' ในเนื้อหาที่เขียน เพราะระบบจะสร้างสารบัญแยกให้อัตโนมัติ
[สำคัญ] ส่งกลับเฉพาะเนื้อหาเรื่องเท่านั้น ห้ามใส่ชื่อเรื่อง (title) ชื่อบท (chapter list) หรือโครงสร้างหนังสือซ้ำอีก
"""


class AutomatedNovelGenerator:
    """Pipeline for generating a full novel via Gemini API."""

    def __init__(self, genre: str, premise: str, total_chapters: int,
                 language: str = "thai", progress_cb: Optional[Callable] = None):
        self.genre = genre
        self.premise = premise
        self.total_chapters = total_chapters
        self.language = language
        self.progress_cb = progress_cb or (lambda msg: None)

        self.client = genai.Client()
        self.model_id = "gemini-3.1-pro-preview"

        self.outline: Optional[NovelOutline] = None
        self.prologue_content: str = ""
        self.chapters_content: List[str] = []

        # Temp dir for caching intermediate results
        self._cache_dir = tempfile.mkdtemp(prefix="novel_cache_")

    @staticmethod
    def _safe_filename(title: str, max_chars: int = 30) -> str:
        """Create a filesystem-safe title for filenames (truncated, no illegal chars)."""
        clean = re.sub(r'[\\/:*?"<>|]', '', title).strip()
        if len(clean) > max_chars:
            cut = clean[:max_chars].rfind(' ')
            clean = clean[:cut].strip() if cut > 0 else clean[:max_chars]
        return clean or 'Novel'

    def _log(self, msg: str):
        logger.info(msg)
        self.progress_cb(msg)

    def _sanitize_content(self, text: str) -> str:
        """Strip accidental TOC, chapter lists, and markdown wrappers from generated text."""
        lines = text.split('\n')
        cleaned = []
        skip_mode = False
        for line in lines:
            stripped = line.strip()
            # Detect TOC header — skip it and subsequent chapter listing lines
            if re.match(r'^(สารบัญ|Table of Contents|TOC)$', stripped, re.IGNORECASE):
                skip_mode = True
                continue
            # While in skip mode, skip lines that look like chapter listings
            if skip_mode:
                if re.match(r'^(บทที่\s*\d|บทนำ|Chapter\s*\d|Prologue|Epilogue)', stripped, re.IGNORECASE):
                    continue
                elif stripped == '' or re.match(r'^(#{1,3}\s|---)', stripped):
                    continue
                else:
                    skip_mode = False  # Content resumed
            cleaned.append(line)
        return '\n'.join(cleaned).strip()

    def _call_gemini(self, prompt: str, temperature: float = 0.8,
                     schema=None, max_retries: int = 3) -> str:
        """Call Gemini API with auto-retry logic."""
        config_kwargs = {"temperature": temperature}
        if schema:
            config_kwargs["response_mime_type"] = "application/json"
            config_kwargs["response_schema"] = schema

        for attempt in range(max_retries):
            try:
                response = self.client.models.generate_content(
                    model=self.model_id,
                    contents=[MASTER_PROMPT, prompt],
                    config=types.GenerateContentConfig(**config_kwargs),
                )
                text = response.text.strip() if response.text else ""
                if text:
                    return text
            except Exception as e:
                logger.warning("Gemini call failed (attempt %d/%d): %s",
                               attempt + 1, max_retries, e)
                time.sleep(2 * (attempt + 1))

        raise RuntimeError(f"Gemini API failed after {max_retries} retries")

    def generate_outline(self):
        """Step 1: Generate novel outline with characters and 3-act structure."""
        self._log("📋 ขั้นตอนที่ 1: กำลังสร้างโครงเรื่องและตัวละคร...")

        lang_instruction = "เขียนทุกอย่างเป็นภาษาไทย" if self.language == "thai" else "Write everything in English"

        prompt = f"""
        {lang_instruction}
        จากประเภทนิยาย '{self.genre}' และพล็อตเรื่อง '{self.premise}'
        สร้างโครงเรื่องนิยายอย่างละเอียดสำหรับ {self.total_chapters} บท
        รวมถึงข้อมูลตัวละครที่มีปมทางจิตใจและเป้าหมายภายนอก
        วางโครงเรื่องตามโครงสร้าง 3 องก์ (3-Act Structure)

        [สำคัญ: การตั้งชื่อเรื่อง]
        ตั้งชื่อเรื่อง (title) ตามกฎเหล่านี้:
        1. สั้น 1-5 คำ เท่านั้น — ห้ามเกิน 5 คำ
        2. สร้าง Curiosity Gap — ชื่อต้องทำให้ผู้อ่านอยากรู้ว่า "เกิดอะไรขึ้น?"
        3. บ่งบอกแนวเรื่อง — ชื่อต้อง hint ว่าเรื่องเป็นแนวอะไร
        4. ห้ามใช้คำซ้ำซากที่ AI มักใช้: shadows, echoes, whispers, fate, destiny, ashes, fragments, silent, hidden, broken, เงา, สะท้อน, กระซิบ, โชคชะตา, เถ้าถ่าน, เศษเสี้ยว
        5. ใช้เทคนิคใดเทคนิคหนึ่ง:
           - Emotional Tension: ชื่อที่สร้างอารมณ์ขัดแย้ง (เช่น "สิ่งที่ไม่เคยบอก", "คนแปลกหน้าข้างกาย")
           - Pattern Interruption: ชื่อที่ทำให้หยุดคิด (เช่น "เราเคยโกหก", "วันที่ฝนหยุดตก")
           - Concrete Image: ชื่อที่วาดภาพชัดเจน (เช่น "ร้านกาแฟปลายสาย", "สวนลับหลังบ้าน")
        6. ทดสอบในใจ: ถ้าชื่อนี้ฟังเหมือนนิยายอื่น 10 เรื่อง → เปลี่ยนใหม่
        """

        text = self._call_gemini(prompt, temperature=0.7, schema=NovelOutline)
        self.outline = NovelOutline.model_validate_json(text)
        self._log(f"✅ สร้างโครงเรื่องสำเร็จ — \"{ self.outline.title}\" | {len(self.outline.characters)} ตัวละคร, {len(self.outline.acts)} องก์")

    def _critique_text(self, text: str) -> str:
        """Run Seven Sweeps auto-critique editor loop with AI-ism detection."""
        critique_prompt = f"""
        คุณคือบรรณาธิการมืออาชีพระดับโลก ตรวจสอบข้อความตามกรอบ Seven Sweeps:

        [Sweep 1: Clarity] ทุกประโยคเข้าใจได้ทันที ไม่กำกวม สรรพนามชัดเจน
        [Sweep 2: Voice] ระดับภาษาสม่ำเสมอ เสียงตัวละครไม่ปนเสียงผู้เขียน
        [Sweep 3: So What] ทุกคำบรรยายต้องรับใช้เรื่องราว ถ้าไม่มีประโยชน์ให้ตัดออก
        [Sweep 4: Prove It] Show Don't Tell — แสดงอารมณ์ผ่านการกระทำและประสาทสัมผัส ห้ามสรุปอารมณ์ตรงๆ
        [Sweep 5: Specificity] แทนที่คำคลุมเครือด้วยรายละเอียดเฉพาะ เพิ่มประสาทสัมผัส 5
        [Sweep 6: Emotion] อารมณ์ต้องเกิดจากเหตุการณ์ ไม่ใช่จากการป้อนความรู้สึกให้ผู้อ่าน
        [Sweep 7: Flow] สลับความยาวประโยค ประโยคเปิดย่อหน้าต้องแข็งแรง ไหลลื่นระหว่างย่อหน้า

        [AI-ISM AUDIT — ตรวจจับและแก้ไขทันที]
        ลบหรือแทนที่คำเหล่านี้:
        - leverage→ใช้, utilize→ใช้, robust→แข็งแกร่ง, seamless→ราบรื่น, cutting-edge→ทันสมัย
        - facilitate→ช่วย, streamline→ทำให้ง่ายขึ้น, navigate→ฝ่าฟัน, pivotal→สำคัญ
        - embark→เริ่มต้น, foster→สร้าง, harness→ใช้, paradigm→แนวทาง
        - tapestry→ส่วนผสม, symphony→การผสาน, testament→พิสูจน์
        ลบรูปแบบเหล่านี้:
        - "Moreover", "Furthermore", "Additionally" (transition filler)
        - "It's worth noting", "It should be noted" (hedging)
        - "Very", "Really", "Extremely", "Incredibly" (hollow intensifiers)
        - ห้ามมีประโยคเริ่มด้วย "It is" หรือ "There are" เกิน 1 ต่อย่อหน้า

        [CHARACTER CONSISTENCY]
        ตรวจสอบความสอดคล้องกับคัมภีร์ตัวละคร: {self.outline.model_dump_json() if self.outline else ''}

        [WORD-LEVEL QUICK FIXES]
        ลบคำเหล่านี้ทุกครั้งที่เจอ: just, actually, basically, really, very, quite, rather
        แทน "in order to" → "เพื่อ" | แทน "things" → คำเฉพาะเจาะจง

        ส่งกลับ **เฉพาะ** ข้อความที่ขัดเกลาแล้ว ห้ามใส่คำอธิบาย ห้ามใส่หมายเหตุ

        ข้อความ:
        {text}
        """
        try:
            return self._call_gemini(critique_prompt, temperature=0.7)
        except Exception:
            return text  # Fallback to unpolished

    def generate_prologue(self):
        """Step 2: Generate the prologue."""
        self._log("📖 ขั้นตอนที่ 2: กำลังเขียนบทนำ (Prologue)...")

        cache_file = os.path.join(self._cache_dir, "prologue.txt")
        if os.path.exists(cache_file):
            with open(cache_file, "r", encoding="utf-8") as f:
                self.prologue_content = f.read()
            self._log("📖 พบแคชบทนำ — ข้ามการสร้างใหม่")
            return

        prompt = f"""
        เขียนบทนำ (Prologue) ที่น่าติดตามสำหรับนิยายเรื่องนี้
        คัมภีร์เรื่อง: {self.outline.model_dump_json()}
        ประเภท: {self.outline.genre}
        พล็อตเรื่อง: {self.outline.premise}

        กฎ:
        - ดึงดูดผู้อ่านทันที กระชับ ทรงพลัง
        - ส่งกลับ **เฉพาะ** เนื้อหาบทนำเท่านั้น
        - **ห้าม** ใส่สารบัญ, ห้ามใส่รายชื่อบท, ห้ามใส่ชื่อเรื่อง, ห้ามใส่ Table of Contents
        - **ห้าม** ครอบด้วย markdown (```, ##, ** เป็นต้น)
        - เริ่มด้วยเนื้อเรื่องทันที
        """
        raw = self._call_gemini(prompt)
        critiqued = self._critique_text(raw)
        self.prologue_content = self._sanitize_content(critiqued)

        with open(cache_file, "w", encoding="utf-8") as f:
            f.write(self.prologue_content)
        self._log("✅ บทนำเสร็จสมบูรณ์")

    def generate_chapters(self):
        """Step 3: Generate all chapters scene-by-scene."""
        self._log(f"✍️ ขั้นตอนที่ 3: กำลังเขียนทั้ง {self.total_chapters} บท...")

        for ch_num in range(1, self.total_chapters + 1):
            self.generate_chapter(ch_num)

    def generate_chapter(self, ch_num: int):
        """Generate a specific chapter."""
        self._log(f"--- กำลังเขียนบทที่ {ch_num}/{self.total_chapters} ---")

        cache_file = os.path.join(self._cache_dir, f"chapter_{ch_num}.txt")
        if os.path.exists(cache_file):
            with open(cache_file, "r", encoding="utf-8") as f:
                self.chapters_content.append(f.read())
            self._log(f"📄 พบแคชบทที่ {ch_num} — ข้ามการสร้างใหม่")
            return

        # 3a. Generate chapter outline (scenes)
        prev_summary = "ยังไม่มีเนื้อหาก่อนหน้า"
        if self.chapters_content:
            prev_summary = self.chapters_content[-1][-500:]

        outline_prompt = f"""
        สร้างโครงร่างสำหรับบทที่ {ch_num} จากทั้งหมด {self.total_chapters} บท
        คัมภีร์เรื่อง: {self.outline.model_dump_json()}
        บริบทก่อนหน้า: {prev_summary}
        ให้มี 3-4 ฉาก พร้อมชื่อบทที่น่าดึงดูดใจ
        """
        ch_outline_json = self._call_gemini(outline_prompt, temperature=0.7,
                                             schema=ChapterOutline)
        ch_outline = ChapterOutline.model_validate_json(ch_outline_json)
        self._log(f"📝 โครงร่างบทที่ {ch_num}: '{ch_outline.title}' ({len(ch_outline.scenes)} ฉาก)")

        # 3b. Generate each scene
        full_chapter = f"บทที่ {ch_num}: {ch_outline.title}\n\n"
        scene_context = ""

        for scene in ch_outline.scenes:
            self._log(f"   ✏️ เขียนฉากที่ {scene.scene_number}: {scene.action_or_conflict[:50]}...")

            scene_prompt = f"""
            เขียนฉากที่ {scene.scene_number} สำหรับบทที่ {ch_num}
            คัมภีร์เรื่อง: {self.outline.model_dump_json()}
            ชื่อบท: {ch_outline.title}
            Setting: {scene.setting}
            Action/Conflict: {scene.action_or_conflict}
            บริบทฉากก่อนหน้า: {scene_context}

            กฎ: Show Don't Tell, เน้นประสาทสัมผัส 5, ใช้คำกริยาทรงพลัง
            ส่งกลับเฉพาะข้อความฉาก ห้ามครอบด้วย markdown
            """
            raw_scene = self._call_gemini(scene_prompt)
            polished_scene = self._critique_text(raw_scene)

            full_chapter += polished_scene + "\n\n***\n\n"
            scene_context = polished_scene[-500:]

        full_chapter = full_chapter.strip().rstrip('*').strip()
        self.chapters_content.append(full_chapter)

        with open(cache_file, "w", encoding="utf-8") as f:
            f.write(full_chapter)
        self._log(f"✅ บทที่ {ch_num} เสร็จสมบูรณ์")

    def export_to_docx(self, title: str = "Novel") -> str:
        """Step 4: Compile everything into a .docx file. Returns filepath."""
        self._log("📦 ขั้นตอนที่ 4: กำลังรวบรวมเป็นไฟล์ .docx...")

        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # --- Set Thai-compatible font for Normal style ---
        THAI_FONT = 'Sarabun'
        style = doc.styles['Normal']
        style.font.name = THAI_FONT
        style.font.size = Pt(14)
        # Set complex-script font (critical for Thai vowel/tone mark rendering)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), THAI_FONT)
        rFonts.set(qn('w:hAnsi'), THAI_FONT)
        rFonts.set(qn('w:cs'), THAI_FONT)
        # Enable complex script for Thai
        cs_elem = rPr.find(qn('w:cs'))
        if cs_elem is None:
            cs_elem = OxmlElement('w:cs')
            rPr.append(cs_elem)

        def _set_heading_font(heading_para):
            """Apply Thai font to all runs in a heading paragraph."""
            for run in heading_para.runs:
                run.font.name = THAI_FONT
                rp = run._element.get_or_add_rPr()
                rf = rp.find(qn('w:rFonts'))
                if rf is None:
                    rf = OxmlElement('w:rFonts')
                    rp.insert(0, rf)
                rf.set(qn('w:ascii'), THAI_FONT)
                rf.set(qn('w:hAnsi'), THAI_FONT)
                rf.set(qn('w:cs'), THAI_FONT)

        # Title page
        t = doc.add_heading(title, 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_heading_font(t)
        doc.add_page_break()

        # Table of Contents (auto-generated by code, NOT by LLM)
        toc = doc.add_heading("สารบัญ", level=1)
        toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_heading_font(toc)
        if self.prologue_content:
            doc.add_paragraph("บทนำ (Prologue)")
        for i, ch in enumerate(self.chapters_content, 1):
            first_line = ch.strip().split('\n')[0] if ch.strip() else f"บทที่ {i}"
            doc.add_paragraph(first_line)
        doc.add_page_break()

        # Prologue (sanitized to remove any accidental TOC from LLM)
        if self.prologue_content:
            h = doc.add_heading("บทนำ (Prologue)", level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_heading_font(h)
            clean_prologue = self._sanitize_content(self.prologue_content)
            for para in clean_prologue.split('\n'):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.first_line_indent = Pt(18)
            doc.add_page_break()

        # Chapters (sanitized)
        for i, ch in enumerate(self.chapters_content, 1):
            clean_ch = self._sanitize_content(ch)
            lines = clean_ch.strip().split('\n')
            ch_title = lines[0].strip() if lines else f"บทที่ {i}"
            content_lines = lines[1:] if lines else []

            h = doc.add_heading(ch_title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_heading_font(h)

            for para in content_lines:
                if para.strip() and para.strip() != '***':
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.first_line_indent = Pt(18)
            doc.add_page_break()

        # Save
        file_id = str(uuid.uuid4())
        short_title = self._safe_filename(title)
        filename = f"{short_title}_{file_id[:8]}.docx"
        filepath = os.path.join(EXPORT_DIR, filename)
        
        try:
            doc.save(filepath)
        except Exception as e:
            self._log(f"⚠️ ไม่สามารถบันทึกไฟล์ด้วยชื่อปกติได้ ลองใช้ชื่อแบบสุ่ม: {e}")
            filename = f"novel_fallback_{file_id[:8]}.docx"
            filepath = os.path.join(EXPORT_DIR, filename)
            doc.save(filepath)

        self._log(f"🎉 เขียนหนังสือเสร็จสิ้น — ไฟล์: {filename}")
        return filename

    def export_to_pdf(self, title: str = "Novel") -> str:
        """Step 5: Generate a PDF version of the novel. Returns the pdf filename."""
        self._log("📄 กำลังสร้างไฟล์ PDF...")

        from weasyprint import HTML

        # Font paths
        fonts_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "fonts")
        font_regular = os.path.join(fonts_dir, "Sarabun-Regular.ttf")
        font_bold = os.path.join(fonts_dir, "Sarabun-Bold.ttf")

        # Build HTML content
        chapters_html = ""

        # Prologue (sanitized)
        if self.prologue_content:
            clean_prologue = self._sanitize_content(self.prologue_content)
            prologue_paras = "\n".join(
                f"<p>{line.strip()}</p>" for line in clean_prologue.split('\n') if line.strip()
            )
            chapters_html += f'<h2 style="text-align:center; page-break-before:always;">บทนำ (Prologue)</h2>\n{prologue_paras}\n'

        # Chapters (sanitized)
        for i, ch in enumerate(self.chapters_content, 1):
            clean_ch = self._sanitize_content(ch)
            lines = clean_ch.strip().split('\n')
            ch_title = lines[0].strip() if lines else f"บทที่ {i}"
            content_lines = lines[1:] if lines else []
            body_paras = "\n".join(
                f"<p>{line.strip()}</p>" for line in content_lines
                if line.strip() and line.strip() != '***'
            )
            chapters_html += f'<h2 style="text-align:center; page-break-before:always;">{ch_title}</h2>\n{body_paras}\n'

        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
    @font-face {{
        font-family: 'Sarabun';
        src: url('file://{font_regular}') format('truetype');
        font-weight: normal;
    }}
    @font-face {{
        font-family: 'Sarabun';
        src: url('file://{font_bold}') format('truetype');
        font-weight: bold;
    }}
    body {{
        font-family: 'Sarabun', sans-serif;
        font-size: 14pt;
        line-height: 1.8;
        margin: 2.5cm;
        color: #1a1a1a;
    }}
    h1 {{
        font-size: 28pt;
        text-align: center;
        margin-top: 40%;
        color: #2c3e50;
    }}
    h2 {{
        font-size: 20pt;
        font-weight: bold;
        color: #2c3e50;
        margin-bottom: 20px;
    }}
    p {{
        text-indent: 2em;
        margin-bottom: 8px;
        word-break: break-word;
    }}
</style>
</head>
<body>
    <h1>{title}</h1>
    {chapters_html}
</body></html>"""

        file_id = str(uuid.uuid4())
        short_title = self._safe_filename(title)
        filename = f"{short_title}_{file_id[:8]}.pdf"
        filepath = os.path.join(EXPORT_DIR, filename)

        try:
            HTML(string=full_html).write_pdf(filepath)
        except Exception as e:
            self._log(f"⚠️ ไม่สามารถบันทึกไฟล์ด้วยชื่อปกติได้ ลองใช้ชื่อแบบสุ่ม: {e}")
            filename = f"novel_fallback_{file_id[:8]}.pdf"
            filepath = os.path.join(EXPORT_DIR, filename)
            HTML(string=full_html).write_pdf(filepath)
            
        self._log(f"📄 สร้าง PDF เสร็จสิ้น — ไฟล์: {filename}")
        return filename

    def run_full_pipeline(self, title: str = "Novel") -> dict:
        """Run the complete novel generation pipeline. Returns dict with filenames."""
        self.generate_outline()
        # Prefer AI-generated title from outline (designed with bestseller psychology)
        # over the raw premise that was passed in as title
        final_title = (self.outline.title if self.outline and self.outline.title else title)
        self._log(f"📚 ชื่อหนังสือ: \"{final_title}\"")
        self.generate_prologue()
        self.generate_chapters()
        docx_file = self.export_to_docx(final_title)
        pdf_file = self.export_to_pdf(final_title)
        return {"docx": docx_file, "pdf": pdf_file, "title": final_title}
