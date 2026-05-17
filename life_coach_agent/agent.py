"""
Life Coach Agent — ADK Agent Definition
Specializes in: Writing coaching, Self-discovery, Life guidance, MCQ Generation,
                Document Export (DOCX, PDF, XLSX)
"""

import os
import uuid
import json
import re
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from duckduckgo_search import DDGS
import requests
from bs4 import BeautifulSoup
import markdownify
from google import genai

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from openpyxl import Workbook
from openpyxl.styles import Font as XlFont, Alignment, PatternFill, Border, Side

from google.adk.agents import Agent

# Directory to store generated files
EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)

logger = logging.getLogger(__name__)

# --- Try to register a Thai-friendly font for PDF ---
_THAI_FONT_NAME = "Helvetica"  # fallback
_THAI_FONT_REGISTERED = False

# Common Thai font paths on macOS
_THAI_FONT_CANDIDATES = [
    "/Library/Fonts/Thonburi.ttc",
    "/System/Library/Fonts/Thonburi.ttc",
    "/Library/Fonts/TH Sarabun New.ttf",
    "/System/Library/Fonts/Supplemental/Thonburi.ttf",
]

for _fpath in _THAI_FONT_CANDIDATES:
    if os.path.exists(_fpath):
        try:
            pdfmetrics.registerFont(TTFont("ThaiFont", _fpath, subfontIndex=0))
            _THAI_FONT_NAME = "ThaiFont"
            _THAI_FONT_REGISTERED = True
            break
        except Exception:
            continue


SYSTEM_INSTRUCTION = """คุณคือ "Life Coach AI" — ผู้ช่วยอัจฉริยะที่เชี่ยวชาญ 3 ด้านหลัก:

## 0. การทำความรู้จักผู้ใช้ (Initial Profiling Survey)
- **เมื่อเริ่มต้นสนทนาครั้งแรก หรือยังไม่มีข้อมูล:** ให้เริ่มด้วยการทำความรู้จักผู้ใช้ก่อนเสมอ โดยใช้รูปแบบ "survey" JSON format (ดูในหัวข้อ 4.1)
- **มิติข้อมูลที่ต้องการดึงจากผู้ใช้ (อ้างอิงจากบทวิเคราะห์จิตวิทยา):**
  1. ค่านิยมและแรงจูงใจ (Values & Motivation)
  2. สไตล์การทำงานและการใช้ชีวิต (Working Style & Preferences)
  3. ความท้าทายหรืออุปสรรคที่ฉุดรั้ง (Challenges & Bottlenecks)
  4. สไตล์การโค้ชชิ่งที่ผู้ใช้ต้องการ (Preferred Coaching Style)
- ปรับน้ำเสียงให้เป็นกันเอง ไม่คุกคาม และทำให้ผู้ใช้รู้สึกปลอดภัยที่จะแชร์ข้อมูลอย่างเปิดเผย

## 1. สอนเขียนหนังสือ (Writing Coach)
- ช่วยพัฒนาทักษะการเขียน ตั้งแต่โครงสร้าง ไปจนถึงสำนวน
- ให้ feedback ที่สร้างสรรค์ วิเคราะห์จุดแข็ง/จุดอ่อนของงานเขียน
- แนะนำเทคนิคการเขียนแบบมืออาชีพ

## 2. ค้นหาตัวตน (Self-Discovery)
- ถามคำถามที่กระตุ้นความคิดเพื่อช่วยผู้ใช้ค้นหาตัวเอง
- วิเคราะห์ค่านิยม จุดแข็ง ความสนใจ
- เสนอมุมมองใหม่ๆ ในการมองตัวเอง

## 3. แนวทางชีวิต (Life Guidance)
- ให้คำปรึกษาเรื่องการวางแผนชีวิต
- ช่วยตั้งเป้าหมายและวางแผนการดำเนินงาน
- ให้กำลังใจและแรงบันดาลใจ

## 4. รูปแบบการสร้างข้อสอบและแบบสอบถาม (JSON Forms)
**สำคัญมาก:** เมื่อสร้าง JSON ให้ตอบกลับเฉพาะ JSON ด้านล่างเท่านั้น ห้ามเพิ่มข้อความอื่นก่อนหรือหลัง JSON

### 4.1 สร้างแบบสอบถาม (Survey Generation)
ใช้สำหรับการทำความรู้จักผู้ใช้ตั้งแต่เริ่มแรก เพื่อลดระยะเวลาการถามทีละข้อ ให้ **สร้างแบบสอบถาม 3-4 ข้อรวดเดียว** ให้ครอบคลุมทุกมิติ (เช่น อาชีพ, เป้าหมาย, ปัญหา) โครงสร้าง JSON:

```json
{
  "type": "survey",
  "title": "ทำความรู้จักกันก่อนเริ่มโค้ชชิ่ง",
  "questions": [
    {
      "id": "q1",
      "question": "คำถามข้อที่ 1",
      "options": [
        {"key": "A", "text": "ตัวเลือก A"},
        {"key": "B", "text": "ตัวเลือก B"},
        {"key": "Other", "text": "อื่นๆ (โปรดระบุ)"}
      ]
    },
    {
      "id": "q2",
      "question": "คำถามข้อที่ 2",
      "options": [
        {"key": "A", "text": "ตัวเลือก A"},
        {"key": "B", "text": "ตัวเลือก B"},
        {"key": "Other", "text": "อื่นๆ (โปรดระบุ)"}
      ]
    }
  ]
}
```
**บังคับ:** ทุกคำถามต้องมีตัวเลือก "อื่นๆ (โปรดระบุ)" เสมอ โดยใช้ key เป็น "Other" หรือคีย์ที่เหมาะสม

### 4.2 สร้างข้อสอบ Multiple Choice (MCQ Generation)
ใช้สำหรับประเมินความรู้หรือทดสอบผู้ใช้ โดยมีคำตอบที่ถูกเพียงข้อเดียว โครงสร้าง JSON:

```json
{
  "type": "mcq",
  "question": "คำถาม",
  "options": [
    {"key": "A", "text": "ตัวเลือก A"},
    {"key": "B", "text": "ตัวเลือก B"},
    {"key": "C", "text": "ตัวเลือก C"},
    {"key": "D", "text": "ตัวเลือก D"}
  ],
  "correct_answer": "A",
  "explanation": "คำอธิบายว่าทำไมคำตอบนี้ถูกต้อง"
}
```

## 5. สร้างไฟล์เอกสาร (Document Export)
เมื่อผู้ใช้ขอให้สร้างไฟล์ ให้เลือก tool ที่ตรงกับรูปแบบที่ผู้ใช้ต้องการ:

### 5.1 ไฟล์ Word (.docx)
- ใช้ tool `create_docx_document` เมื่อผู้ใช้ขอไฟล์ .docx หรือเอกสาร Word
- เหมาะกับ: รายงาน, บทความ, คู่มือ, เนื้อหายาว

### 5.2 ไฟล์ PDF (.pdf)
- ใช้ tool `create_pdf_document` เมื่อผู้ใช้ขอไฟล์ .pdf
- เหมาะกับ: เอกสารที่ต้องการอ่านอย่างเดียว, รายงาน, สรุป

### 5.3 ไฟล์ Excel (.xlsx)
- ใช้ tool `create_xlsx_document` เมื่อผู้ใช้ขอไฟล์ .xlsx หรือ Excel หรือตาราง
- เหมาะกับ: ข้อมูลตาราง, แผน, checklist, การเปรียบเทียบ
- **สำคัญ:** ให้ส่ง headers และ rows เป็น JSON array
  - headers: ["หัวข้อ1", "หัวข้อ2", ...]
  - rows: [["ข้อมูล1", "ข้อมูล2", ...], [...], ...]

### หลักการเลือกรูปแบบ:
- ถ้าผู้ใช้ระบุชัดเจน (เช่น "สร้าง PDF", "ทำ Excel") → ใช้ตามที่ระบุ
- ถ้าผู้ใช้พูดว่า "สร้างไฟล์" หรือ "สร้างเอกสาร" โดยไม่ระบุรูปแบบ → ใช้ docx
- ถ้าเนื้อหาเป็นตาราง/ข้อมูล → แนะนำ xlsx
- หลังจากเรียก tool สำเร็จ ให้ตอบกลับเฉพาะ JSON ที่ tool ส่งกลับมา ห้ามเพิ่มข้อความอื่น

## 6. เครื่องมืออินเทอร์เน็ตและการสร้างภาพ (Web Tools & Image Generation)
- **การค้นหาข้อมูลบนเว็บ (Web Search):** หากผู้ใช้สอบถามข้อมูลที่เป็นปัจจุบัน ข้อมูลข่าวสาร หรือข้อมูลที่คุณไม่แน่ใจ ให้ใช้ tool `web_search` เพื่อค้นหาข้อมูลจากอินเทอร์เน็ต
- **การอ่านเนื้อหาจากลิงก์ (Web Scraping):** หากผู้ใช้ส่ง URL มาให้ หรือต้องการให้สรุปข้อมูลจากเว็บไซต์ใดๆ ให้ใช้ tool `read_url` เพื่อดึงเนื้อหามาอ่านและวิเคราะห์
- **การสร้างภาพ (Image Generation):** หากผู้ใช้ขอให้สร้างภาพ วาดรูป หรือจินตภาพเป็นรูปภาพ ให้ใช้ tool `generate_image` โดยให้ระบุ prompt อธิบายภาพอย่างละเอียดเป็นภาษาอังกฤษ (คุณสามารถแปลจากภาษาไทยของผู้ใช้ให้เป็น prompt ภาษาอังกฤษที่สละสลวยได้) หลังจากเรียกสำเร็จ ให้ส่ง JSON ที่ได้รับกลับไปให้ผู้ใช้

## แนวทางการสื่อสาร
- ใช้ภาษาไทยเป็นหลัก แต่ถ้าผู้ใช้พูดภาษาอังกฤษ ให้ตอบเป็นภาษาอังกฤษ
- พูดเป็นกันเอง อบอุ่น แต่ให้ข้อมูลที่มีคุณภาพ
- ถามคำถาม follow-up เพื่อเข้าใจผู้ใช้มากขึ้น
- ให้ตัวอย่างประกอบเมื่อเป็นไปได้
- **สำคัญ:** เมื่อผู้ใช้ขอสร้างไฟล์หรือสร้างภาพ ให้เรียกใช้ tool ที่เหมาะสมทุกครั้ง
"""


# --- Helper: parse markdown-ish content into lines ---
def _parse_content_lines(content: str):
    """Parse content string into structured line tuples: (type, text)."""
    lines = content.split('\n')
    parsed = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            parsed.append(("blank", ""))
        elif stripped.startswith('## '):
            parsed.append(("h2", stripped[3:].strip()))
        elif stripped.startswith('### '):
            parsed.append(("h3", stripped[4:].strip()))
        elif stripped.startswith('- ') or stripped.startswith('• '):
            parsed.append(("bullet", stripped[2:].strip()))
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s*', '', stripped)
            parsed.append(("numbered", text))
        else:
            parsed.append(("text", stripped))
    return parsed


def generate_mcq(topic: str, difficulty: str = "medium") -> dict:
    """Generate a multiple choice question on a given topic.

    Args:
        topic: The topic for the question (e.g., "writing techniques", "life planning")
        difficulty: Difficulty level — "easy", "medium", or "hard"

    Returns:
        dict: A status dict instructing the agent to produce a structured MCQ JSON.
    """
    return {
        "status": "success",
        "instruction": f"Please generate a multiple choice question about '{topic}' "
                       f"at {difficulty} difficulty level. "
                       f"Return ONLY the JSON format as specified in your system instructions."
    }


def create_docx_document(title: str, content: str) -> dict:
    """Create a .docx document file that the user can download.

    Use this tool whenever the user asks to create a Word document or .docx file.

    Args:
        title: The document title (e.g., "คู่มือการสร้างหนังสือ", "แผนพัฒนาตนเอง")
        content: The full document content. Use '\\n' for line breaks and '## ' prefix for section headings.

    Returns:
        dict: A JSON object with type "docx_download" containing the file_id and filename for download.
    """
    doc = Document()

    # --- Style the document ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(14)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacer

    # Parse content into sections
    for line_type, text in _parse_content_lines(content):
        if line_type == "blank":
            doc.add_paragraph("")
        elif line_type == "h2":
            doc.add_heading(text, level=2)
        elif line_type == "h3":
            doc.add_heading(text, level=3)
        elif line_type == "bullet":
            doc.add_paragraph(text, style='List Bullet')
        elif line_type == "numbered":
            doc.add_paragraph(text, style='List Number')
        else:
            doc.add_paragraph(text)

    # Save file
    file_id = str(uuid.uuid4())
    safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
    filename = f"{safe_title}_{file_id[:8]}.docx"
    filepath = os.path.join(EXPORT_DIR, filename)
    doc.save(filepath)

    result = {
        "type": "docx_download",
        "file_id": file_id,
        "filename": filename,
        "title": title,
        "message": f"สร้างเอกสาร Word '{title}' เรียบร้อยแล้วครับ! คลิกปุ่มด้านล่างเพื่อดาวน์โหลด 📄"
    }

    return {"status": "success", "result": json.dumps(result, ensure_ascii=False)}


def create_pdf_document(title: str, content: str) -> dict:
    """Create a .pdf document file that the user can download.

    Use this tool whenever the user asks to create a PDF file.

    Args:
        title: The document title (e.g., "สรุปแนวทางการเขียน", "รายงานการค้นหาตัวตน")
        content: The full document content. Use '\\n' for line breaks and '## ' prefix for section headings.

    Returns:
        dict: A JSON object with type "pdf_download" containing the file_id and filename for download.
    """
    file_id = str(uuid.uuid4())
    safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
    filename = f"{safe_title}_{file_id[:8]}.pdf"
    filepath = os.path.join(EXPORT_DIR, filename)

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        leftMargin=25 * mm,
        rightMargin=25 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    # --- Styles ---
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontName=_THAI_FONT_NAME,
        fontSize=22,
        spaceAfter=12,
        textColor=HexColor("#1a1a2e"),
        alignment=1,  # center
    )
    heading2_style = ParagraphStyle(
        'CustomH2',
        parent=styles['Heading2'],
        fontName=_THAI_FONT_NAME,
        fontSize=16,
        spaceBefore=14,
        spaceAfter=6,
        textColor=HexColor("#6c3ce0"),
    )
    heading3_style = ParagraphStyle(
        'CustomH3',
        parent=styles['Heading3'],
        fontName=_THAI_FONT_NAME,
        fontSize=13,
        spaceBefore=10,
        spaceAfter=4,
        textColor=HexColor("#4a4a6a"),
    )
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontName=_THAI_FONT_NAME,
        fontSize=11,
        leading=16,
        spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=body_style,
        leftIndent=20,
        bulletIndent=8,
        spaceAfter=3,
    )

    # --- Build PDF elements ---
    elements = []
    elements.append(Paragraph(title, title_style))
    elements.append(Spacer(1, 10))

    for line_type, text in _parse_content_lines(content):
        if line_type == "blank":
            elements.append(Spacer(1, 6))
        elif line_type == "h2":
            elements.append(Paragraph(text, heading2_style))
        elif line_type == "h3":
            elements.append(Paragraph(text, heading3_style))
        elif line_type == "bullet":
            elements.append(Paragraph(f"• {text}", bullet_style))
        elif line_type == "numbered":
            elements.append(Paragraph(f"   {text}", bullet_style))
        else:
            elements.append(Paragraph(text, body_style))

    doc.build(elements)

    result = {
        "type": "pdf_download",
        "file_id": file_id,
        "filename": filename,
        "title": title,
        "message": f"สร้างไฟล์ PDF '{title}' เรียบร้อยแล้วครับ! คลิกปุ่มด้านล่างเพื่อดาวน์โหลด 📕"
    }

    return {"status": "success", "result": json.dumps(result, ensure_ascii=False)}


def create_xlsx_document(title: str, headers_json: str, rows_json: str) -> dict:
    """Create an .xlsx Excel spreadsheet file that the user can download.

    Use this tool whenever the user asks to create an Excel file, spreadsheet,
    table, checklist, or structured data export.

    Args:
        title: The spreadsheet title (e.g., "แผนการอ่าน 30 วัน", "Checklist พัฒนาตนเอง")
        headers_json: JSON array string of column headers, e.g. '["หัวข้อ", "รายละเอียด", "สถานะ"]'
        rows_json: JSON array of arrays string for row data, e.g. '[["อ่านหนังสือ", "30 นาที/วัน", "กำลังทำ"]]'

    Returns:
        dict: A JSON object with type "xlsx_download" containing the file_id and filename for download.
    """
    # Parse JSON string inputs
    try:
        headers = json.loads(headers_json) if isinstance(headers_json, str) else headers_json
    except (json.JSONDecodeError, TypeError):
        headers = ["Column 1"]

    try:
        rows = json.loads(rows_json) if isinstance(rows_json, str) else rows_json
    except (json.JSONDecodeError, TypeError):
        rows = []

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]  # Excel sheet name max 31 chars

    # --- Header styling ---
    header_font = XlFont(name='Arial', size=12, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='6C3CE0', end_color='6C3CE0', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin', color='D4D4D8'),
        right=Side(style='thin', color='D4D4D8'),
        top=Side(style='thin', color='D4D4D8'),
        bottom=Side(style='thin', color='D4D4D8'),
    )

    # --- Write title row ---
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(headers), 1))
    title_cell = ws.cell(row=1, column=1, value=title)
    title_cell.font = XlFont(name='Arial', size=16, bold=True, color='1A1A2E')
    title_cell.alignment = Alignment(horizontal='center', vertical='center')

    # --- Write headers (row 3) ---
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_idx, value=str(header))
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    # --- Write data rows ---
    data_font = XlFont(name='Arial', size=11)
    data_align = Alignment(vertical='center', wrap_text=True)
    alt_fill = PatternFill(start_color='F5F3FF', end_color='F5F3FF', fill_type='solid')

    for row_idx, row_data in enumerate(rows, 4):
        if not isinstance(row_data, (list, tuple)):
            row_data = [row_data]
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=str(value))
            cell.font = data_font
            cell.alignment = data_align
            cell.border = thin_border
            if (row_idx - 4) % 2 == 1:
                cell.fill = alt_fill

    # --- Auto-fit column widths (approximate) ---
    for col_idx in range(1, len(headers) + 1):
        max_len = len(str(headers[col_idx - 1])) if col_idx <= len(headers) else 10
        for row_data in rows:
            if isinstance(row_data, (list, tuple)) and col_idx - 1 < len(row_data):
                max_len = max(max_len, len(str(row_data[col_idx - 1])))
        col_letter = chr(64 + col_idx) if col_idx <= 26 else 'A'
        ws.column_dimensions[col_letter].width = min(max_len + 4, 50)

    # --- Save ---
    file_id = str(uuid.uuid4())
    safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
    filename = f"{safe_title}_{file_id[:8]}.xlsx"
    filepath = os.path.join(EXPORT_DIR, filename)
    wb.save(filepath)

    result = {
        "type": "xlsx_download",
        "file_id": file_id,
        "filename": filename,
        "title": title,
        "row_count": len(rows),
        "message": f"สร้างไฟล์ Excel '{title}' ({len(rows)} แถว) เรียบร้อยแล้วครับ! คลิกปุ่มด้านล่างเพื่อดาวน์โหลด 📊"
    }

    return {"status": "success", "result": json.dumps(result, ensure_ascii=False)}


def web_search(query: str, max_results: int = 5) -> dict:
    """Perform a web search using DuckDuckGo to find information.
    
    Args:
        query: The search query.
        max_results: Maximum number of results to return (default 5).
        
    Returns:
        dict: A JSON object containing the search results (title, url, body) or error.
    """
    try:
        results = []
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=max_results):
                results.append(r)
        return {"status": "success", "results": results}
    except Exception as e:
        logger.error("web_search failed for query '%s': %s", query, e, exc_info=True)
        return {"status": "error", "message": "ไม่สามารถค้นหาข้อมูลได้ในขณะนี้ กรุณาลองใหม่"}


def read_url(url: str) -> dict:
    """Extract and summarize content from a given URL.
    
    Args:
        url: The URL to scrape.
        
    Returns:
        dict: A JSON object containing the markdown content of the page or error.
    """
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
            
        markdown_text = markdownify.markdownify(str(soup), heading_style="ATX").strip()
        
        # Limit content length to avoid exceeding token limits
        if len(markdown_text) > 15000:
            markdown_text = markdown_text[:15000] + "\n\n...[Content truncated]..."
            
        return {"status": "success", "content": markdown_text}
    except Exception as e:
        logger.error("read_url failed for '%s': %s", url, e, exc_info=True)
        return {"status": "error", "message": "ไม่สามารถอ่านเนื้อหาจาก URL นี้ได้ กรุณาตรวจสอบลิงก์"}


def generate_image(prompt: str) -> dict:
    """Generate an image using the Nano Banana Pro (Gemini 3 Pro Image Preview) model.
    
    Args:
        prompt: A detailed description of the image to generate. Please use English for best results.
        
    Returns:
        dict: A JSON object containing the image generation status and the image URL.
    """
    try:
        # We need a dedicated directory for images in static
        image_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "outputs")
        os.makedirs(image_dir, exist_ok=True)
        
        client = genai.Client()
        
        response = client.models.generate_content(
            model='gemini-3-pro-image-preview',
            contents=prompt,
            config=genai.types.GenerateContentConfig(
                response_modalities=["IMAGE"],
            )
        )
        
        if hasattr(response, 'candidates') and response.candidates and response.candidates[0].content.parts:
            for part in response.candidates[0].content.parts:
                data = None
                if getattr(part, "inline_data", None) and hasattr(part.inline_data, "data"):
                    data = part.inline_data.data
                elif getattr(part, "image", None) and hasattr(part.image, "image_bytes"):
                    data = part.image.image_bytes
                    
                if data:
                    file_id = str(uuid.uuid4())[:8]
                    filename = f"generated_{file_id}.png"
                    filepath = os.path.join(image_dir, filename)
                    
                    with open(filepath, "wb") as f:
                        f.write(data)
                    
                    result = {
                        "type": "image_generation",
                        "filename": filename,
                        "url": f"/images/{filename}",
                        "message": f"วาดภาพ '{prompt}' เสร็จเรียบร้อยแล้วครับ! 🎨"
                    }
                    return {"status": "success", "result": json.dumps(result, ensure_ascii=False)}
                    
        return {"status": "error", "message": "ไม่สามารถสร้างภาพได้จาก Model"}
    except Exception as e:
        logger.error("generate_image failed for prompt '%s': %s", prompt, e, exc_info=True)
        return {"status": "error", "message": "ไม่สามารถสร้างภาพได้ในขณะนี้ กรุณาลองใหม่"}


root_agent = Agent(
    name="life_coach_agent",
    model="gemini-3.1-pro-preview",
    description="AI Life Coach that helps with writing, self-discovery, life guidance, web searching, URL reading, image generation, and creates downloadable documents (DOCX, PDF, XLSX).",
    instruction=SYSTEM_INSTRUCTION,
    tools=[generate_mcq, create_docx_document, create_pdf_document, create_xlsx_document, web_search, read_url, generate_image],
)
